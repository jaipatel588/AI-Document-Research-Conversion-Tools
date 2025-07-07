from pathlib import Path
from PIL import Image, UnidentifiedImageError
from docx import Document
from fpdf import FPDF
from pdf2image import convert_from_path
from PyPDF2 import PdfReader
from pytesseract import image_to_string
import uuid
import logging

# Logger setup
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# Directories
BASE_DIR = Path(__file__).resolve().parent.parent.parent
OUTPUT_DIR = BASE_DIR / "app" / "data" / "converted"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def handle_conversion_to_format(file_path: Path, target_format: str) -> Path | None:
    ext = file_path.suffix.lower()

    try:
        if ext in [".jpg", ".jpeg", ".png"]:
            if target_format == "pdf":
                return convert_image(file_path, to_format="pdf")
            elif target_format == "txt":
                return convert_image_to_text(file_path)

        elif ext == ".pdf":
            if target_format == "jpg":
                return convert_pdf_to_images(file_path)[0]
            elif target_format == "txt":
                return convert_pdf_to_txt(file_path)
            elif target_format == "docx":
                return convert_pdf_to_docx(file_path)

        elif ext == ".docx" and target_format == "pdf":
            return convert_docx_to_pdf(file_path)

        elif ext == ".txt" and target_format == "pdf":
            return convert_txt_to_pdf(file_path)

        logger.warning(f"⚠️ Unsupported conversion from {ext} to {target_format}")
        return None

    except Exception as e:
        logger.error(f"❌ Error during conversion from {file_path.name} to {target_format}: {e}")
        return None

def convert_image(file_path: Path, to_format: str) -> Path:
    try:
        with Image.open(file_path) as img:
            output_file = OUTPUT_DIR / f"{file_path.stem}_{uuid.uuid4().hex}.{to_format}"
            rgb_img = img.convert("RGB") if to_format in ["pdf", "jpg"] else img
            rgb_img.save(str(output_file), format=to_format.upper())
            logger.info(f"✅ Image converted to {to_format.upper()}: {output_file.name}")
            return output_file
    except UnidentifiedImageError:
        logger.error(f"❌ Cannot identify image file: {file_path}")
        raise
    except Exception as e:
        logger.error(f"❌ Image conversion failed: {e}")
        raise

def convert_image_to_text(file_path: Path) -> Path:
    try:
        with Image.open(file_path) as img:
            text = image_to_string(img)

        output_file = OUTPUT_DIR / f"{file_path.stem}_{uuid.uuid4().hex}.txt"
        output_file.write_text(text.strip(), encoding="utf-8")
        logger.info(f"✅ OCR text saved: {output_file.name}")
        return output_file
    except Exception as e:
        logger.error(f"❌ Image to text OCR failed: {e}")
        raise

def convert_pdf_to_images(file_path: Path) -> list[Path]:
    try:
        images = convert_from_path(str(file_path))
        output_files = []

        for i, image in enumerate(images):
            out_path = OUTPUT_DIR / f"{file_path.stem}_page{i+1}.jpg"
            image.save(str(out_path), "JPEG")
            output_files.append(out_path)

        logger.info(f"✅ PDF converted to {len(output_files)} image(s)")
        return output_files
    except Exception as e:
        logger.error(f"❌ PDF to image conversion failed: {e}")
        raise

def convert_pdf_to_txt(file_path: Path) -> Path:
    try:
        reader = PdfReader(str(file_path))
        text = "".join(page.extract_text() or "" for page in reader.pages).strip()
        output_file = OUTPUT_DIR / f"{file_path.stem}_{uuid.uuid4().hex}.txt"
        output_file.write_text(text, encoding="utf-8")
        logger.info(f"✅ PDF to TXT success: {output_file.name}")
        return output_file
    except Exception as e:
        logger.error(f"❌ PDF to TXT failed: {e}")
        raise

def convert_pdf_to_docx(file_path: Path) -> Path:
    try:
        reader = PdfReader(str(file_path))
        doc = Document()
        for page in reader.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text.strip())

        output_file = OUTPUT_DIR / f"{file_path.stem}_{uuid.uuid4().hex}.docx"
        doc.save(str(output_file))
        logger.info(f"✅ PDF to DOCX success: {output_file.name}")
        return output_file
    except Exception as e:
        logger.error(f"❌ PDF to DOCX failed: {e}")
        raise

def convert_docx_to_pdf(file_path: Path) -> Path:
    try:
        doc = Document(str(file_path))
        pdf = FPDF(format='A4')
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        pdf.set_left_margin(10)
        pdf.set_right_margin(10)

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                pdf.multi_cell(0, 10, text)

        output_file = OUTPUT_DIR / f"{file_path.stem}_{uuid.uuid4().hex}.pdf"
        pdf.output(str(output_file))
        logger.info(f"✅ DOCX to PDF success: {output_file.name}")
        return output_file
    except Exception as e:
        logger.error(f"❌ DOCX to PDF failed: {e}")
        raise

def convert_txt_to_pdf(file_path: Path) -> Path:
    try:
        with file_path.open("r", encoding="utf-8") as f:
            lines = f.readlines()

        pdf = FPDF(format='A4')
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.set_left_margin(10)
        pdf.set_right_margin(10)

        for line in lines:
            clean = line.strip()
            if clean:
                pdf.multi_cell(0, 10, clean)

        output_file = OUTPUT_DIR / f"{file_path.stem}_{uuid.uuid4().hex}.pdf"
        pdf.output(str(output_file))
        logger.info(f"✅ TXT to PDF success: {output_file.name}")
        return output_file
    except Exception as e:
        logger.error(f"❌ TXT to PDF failed: {e}")
        raise